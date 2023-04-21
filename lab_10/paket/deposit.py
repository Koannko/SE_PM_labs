from docx import Document
def count_dep(S, m, r):
    for i in range(m):
        S *= (100 + r / 12) / 100
    return int(S)

def count_revenue(S, m, r):
    E = count_dep(S, m, r)
    return E - S

def save_dep(S, n, r):
    A = count_dep(S, r, n)
    E = count_revenue(S, n, r)
    document = Document()

    document.add_heading('Вклад', level=0)

    document.add_heading('Ввод', level=1)
    document.add_paragraph(f'Сумма вклада (в рублях): {S}')
    document.add_paragraph(f'Срок вклада (в месяцах): {n}')
    document.add_paragraph(f'Годовая ставка (в процентах): {r}')

    document.add_heading('Вывод', level=1)
    
    document.add_paragraph(f'Доход по вкладу за весь период: {E} рублей')
    document.add_paragraph(f'Размер вклада на конец периода: {A} рублей')

    document.add_page_break()

    document.save('deposit.docx')
    print('DONE')
