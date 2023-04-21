from docx import Document
def month_pay(S, n, r):
    i = r/100/12
    К = (i * (1 + i)**n) / ((1+i)**n-1)
    return int(К * S)

def overpayment(S, n, r):
    A = month_pay(S, r, n)
    return int(A * n * 12 - S)

def save_ins(S, n, r):
    A = month_pay(S, r, n)
    E = overpayment(S, n, r)
    document = Document()

    document.add_heading('Рассрочка', level=0)

    document.add_heading('Ввод', level=1)
    document.add_paragraph(f'Сумма рассрочки (в рублях): {S}')
    document.add_paragraph(f'Срок рассрочки (в месяцах): {n}')
    document.add_paragraph(f'Месячная ставка (в процентах): {r}')

    document.add_heading('Вывод', level=1)
    
    document.add_paragraph(f'Ежемесячный платеж составит: {A} рублей')
    document.add_paragraph(f'Переплата по рассрочке составит: {E} рублей')

    document.add_page_break()

    document.save('installment.docx')
    print('DONE')