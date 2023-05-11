from abc import ABC, abstractmethod
from datetime import date, timedelta
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

import calendar

class BankService(ABC):
    def calculate_interest(self):
        return round(self.debt * self.interest / 1200, 2)
    
    def get_date_of_payment(self):
        days = calendar.monthrange(self.date.year, self.date.month)[1]
        self.date += timedelta(days=days)
        month = calendar.month_abbr[self.date.month]
        return f'{month} {self.date.day}, {self.date.year}'

    def calculate_overpayments(self):
        return self.calculate_amount_of_payment() - self.amount

    def pay_reg_payment(self):
        self.debt -= self.calculate_main_debt()
        self.current_month += 1

        if round(self.debt) == 0:
            self.debt = round(self.debt, 1)
    
    @abstractmethod
    def calculate_month_payment(self):
        pass

    @abstractmethod
    def calculate_amount_of_payment(self):
        pass

    @abstractmethod
    def check_input(self):
        pass

    @abstractmethod
    def activate(self):
        pass

    @abstractmethod
    def save_docx(self):
        pass

class Credit(BankService):
    def __init__(self, amount, interest, period):
        self.amount = amount
        self.interest = interest
        self.period = period
        self.debt = amount
        self.date = date.today()
        self.current_month = 0
        
    def calculate_month_payment(self):
        return self.amount * self.interest / \
                1200 * (1 + self.interest / 1200) ** self.period / \
                ((1 + self.interest / 1200) ** self.period - 1)
    
    def calculate_amount_of_payment(self):
        return self.calculate_month_payment() * self.period
    
    def calculate_main_debt(self):
        return self.calculate_month_payment() - self.calculate_interest()
    
    def check_input(self):
        try:
            self.amount = float(self.amount)
            self.interest = float(self.interest)
            self.period = int(self.period) 
            self.amount <= 0 or self.interest <= 0 or self.period <= 0
        except:    
            return f'Вы ввели недопустимые данные: {self.amount}, {self.interest}, {self.period}. Кредит не может быть рассчитан.'
        return True

    def activate(self):
        if self.check_input() != True:
            return self.check_input()
        compute_list = [['Ежемесячный платёж', round(self.calculate_month_payment(), 2)],
                    ['Переплата по кредиту', round(self.calculate_overpayments(), 2)],
                    ['Сумма кредита', self.amount],
                    ['График платежей', self.period]]
        data_list = []
        while self.current_month < self.period:
            self.pay_reg_payment()
            data_list.append([['№', str(self.current_month).zfill(2)],
                            ['Дата платежа', self.get_date_of_payment()],
                            ['Платёж в месяц', round(self.calculate_month_payment(), 2)],
                            ['Основной долг', round(self.calculate_main_debt(), 2)],
                            ['Проценты', round(self.calculate_interest(), 2)],
                            ['Остаток', round(self.debt, 2)]])     
        return data_list, compute_list
        
    def save_docx(self, data_lst, copy_num):
        document = Document()
        document.add_heading('Кредитный калькулятор', level=0)
        document.add_heading('Ввeденные данные', level=1)
        document.add_paragraph(f'Сумма кредита  {self.amount} ₽')
        document.add_paragraph(f'Процентная ставка  {self.interest} процентов')
        document.add_paragraph(f'Срок кредитования  {self.period} месяцев')
        document.add_heading('Результаты расчёта', level=1)
        document.add_paragraph(f'Ежемесячный платеж  {round(self.calculate_month_payment(), 2)} ₽')
        document.add_paragraph(f'Переплата по кредиту  {round(self.calculate_overpayments(), 2)} ₽')
        document.add_heading('График платежей', level=1)

        table = document.add_table(rows=1, cols=6, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        column_widths = [Cm(1.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0)]
        for idx, width in enumerate(column_widths):
            table.columns[idx].width = width

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Дата платежа'
        hdr_cells[2].text = 'Платёж в месяц'
        hdr_cells[3].text = 'Основной долг'
        hdr_cells[4].text = 'Проценты'
        hdr_cells[5].text = 'Остаток'

        for i, row in enumerate(data_lst):
            table.add_row()
            for j, val in enumerate(row):
                cell = table.cell(i+1, j)
                if j > 1:
                    cell.text = str(val[1]) + ' ₽'
                else:
                    cell.text = str(val[1])

        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        document.add_page_break()
        if copy_num == 0:
            document.save('credit.docx')
        else:
            document.save(f'credit{copy_num}.docx')


# РАССРОЧКА
class Installment(BankService):
    def __init__(self, amount, interest, period):
        self.amount = amount
        self.interest = interest
        self.period = period
        self.debt = amount
        self.date = date.today()
        self.current_month = 0
        
    def calculate_month_payment(self):
        return self.calculate_interest() + self.amount / self.period
    
    def get_days_in_month(self):
        return calendar.monthrange(self.date.year, self.date.month)[1]

    def get_days_in_year(self):
        return 365 + calendar.isleap(self.date.year)

    def calculate_amount_of_payment(self):
        sum = self.amount
        payments = sum
        diff = self.amount / self.period
        for i in range(self.period):
            payments += self.interest / 1200 * sum
            sum -= diff
        return payments
    
    def calculate_main_debt(self):
        return self.amount / self.period
    
    def get_current_debt(self):
        return self.debt - self.calculate_main_debt()

    def check_input(self):
        try:
            self.amount = float(self.amount)
            self.interest = float(self.interest)
            self.period = int(self.period) 
            self.amount <= 0 or self.interest <= 0 or self.period <= 0
        except:    
            return f'Вы ввели недопустимые данные: {self.amount}, {self.interest}, {self.period}. Рассрочка не может быть рассчитана.'
        return True

    def activate(self):
        if self.check_input() != True:
            return self.check_input()
        compute_list = [['Средний платёж в месяц', round(self.calculate_amount_of_payment() / self.period, 2)],
                        ['Переплата', round(self.calculate_overpayments(), 2)]]
        data_list = []
        while self.current_month < self.period:
            data_list.append([['№', str(self.current_month).zfill(2)],
                            ['Дата платежа', self.get_date_of_payment()],
                            ['Платёж в месяц', round(self.calculate_month_payment(), 2)],
                            ['Основной долг', round(self.calculate_main_debt(), 2)],
                            ['Проценты', round(self.calculate_interest(), 2)],
                            ['Остаток', round(self.get_current_debt(), 2)]]) 
            self.pay_reg_payment()
        return data_list, compute_list
        
    def save_docx(self, data_lst, copy_num):
        document = Document()
        document.add_heading('Калькулятор рассрочки', level=0)
        document.add_heading('Ввeденные данные', level=1)
        document.add_paragraph(f'Сумма покупки  {self.amount} ₽')
        document.add_paragraph(f'Процентная ставка  {self.interest} процентов')
        document.add_paragraph(f'Срок рассрочки  {self.period} месяцев')
        document.add_heading('Результаты расчёта', level=1)
        document.add_paragraph(f'Платёж в месяц в среднем  {round(self.calculate_amount_of_payment(), 2)} ₽')
        document.add_paragraph(f'Переплата  {round(self.calculate_overpayments(), 2)} ₽')
        document.add_heading('График платежей', level=1)

        table = document.add_table(rows=1, cols=6, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        column_widths = [Cm(1.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0)]
        for idx, width in enumerate(column_widths):
            table.columns[idx].width = width

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Дата платежа'
        hdr_cells[2].text = 'Платёж в месяц'
        hdr_cells[3].text = 'Основной долг'
        hdr_cells[4].text = 'Проценты'
        hdr_cells[5].text = 'Остаток'

        for i, row in enumerate(data_lst):
            table.add_row()
            for j, val in enumerate(row):
                cell = table.cell(i+1, j)
                if j > 1:
                    cell.text = str(val[1]) + ' ₽'
                else:
                    cell.text = str(val[1])

        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        document.add_page_break()
        if copy_num == 0:
            document.save('installment.docx')
        else:
            document.save(f'installment{copy_num}.docx')






# ВКЛАД
class Deposit(BankService):
    def __init__(self, amount, interest, period, interest_type, additions, amount_of_additions):
        self.amount = amount
        self.interest = interest
        self.period = period
        self.interest_type = interest_type
        self.additions = additions
        self.amount_of_additions = amount_of_additions
        self.date = date.today()
        self.current_month = 0
        self.debt = self.amount
        
    def calculate_month_payment(self):
        return self.calculate_interest()

    def calculate_amount_of_payment(self):
        if self.additions == 'Каждый месяц':
            return float(self.amount_of_additions) * self.period
        elif self.additions == 'Каждый квартал':
            return float(self.amount_of_additions) * (self.current_month // 3)
        elif self.additions == 'Каждый год':
            return float(self.amount_of_additions) * (self.current_month // 12)
        else:
            return 0
    
    def calculate_main_debt(self):
        return self.amount
    
    def calculate_addition(self):
        if self.current_month == 0: return 0

        if self.additions == 'Каждый месяц':
            return float(self.amount_of_additions)
        elif self.additions == 'Каждый квартал' and self.current_month % 3 == 0:
            return float(self.amount_of_additions)
        elif self.additions == 'Каждый год' and self.current_month % 12 == 0:
            return float(self.amount_of_additions)
        else:
            return 0
    
    def pay_reg_payment(self):
        self.amount += float(self.calculate_addition())
        self.current_month += 1
        
        if self.interest_type == True:
            print('YES')
            self.amount += self.calculate_interest()
            
    
    def check_input(self):
        try:
            self.amount = float(self.amount)
            self.interest = float(self.interest)
            self.period = int(self.period) 
            self.amount <= 0 or self.interest <= 0 or self.period <= 0 or self.amount_of_additions == ''
        except:    
            return f'Вы ввели недопустимые данные: {self.amount}, {self.interest}, {self.period}. Вклад не может быть рассчитана.'
        return True

    def activate(self):
        if self.check_input() != True:
            return self.check_input()
        data_list = []
        compute_list = []
        while self.current_month < self.period:
            self.pay_reg_payment()
            lst = [['№', str(self.current_month).zfill(2)],
                            ['Дата начисления', self.get_date_of_payment()],
                            ['Начислено процентов', round(self.calculate_month_payment(), 2)],
                            ['Пополнение вклада', round(self.calculate_addition(), 2)],
                            ['Остаток вклада', round(self.calculate_main_debt(), 2)]]
            if self.current_month == self.period:
                compute_list.append(['Сумма в конце срока', round(lst[4][1], 2)])
                compute_list.append(['Доход', round(lst[4][1] - self.debt, 2)])
            data_list.append(lst) 
        print(compute_list)   
        return data_list, compute_list
        
    def save_docx(self, data_lst, compute_list, copy_num):
        if self.interest_type == True: interest_type = 'Да' 
        else: interest_type = 'Нет'
        document = Document()
        document.add_heading('Калькулятор доходности вкладов', level=0)
        document.add_heading('Ввeденные данные', level=1)
        document.add_paragraph(f'Сумма вклада  {self.debt} ₽')
        document.add_paragraph(f'Процентная ставка  {self.interest} процентов')
        document.add_paragraph(f'Срок размещения  {self.period} месяцев')
        document.add_paragraph(f'Добавлять начисления ко вкладу  {interest_type}')
        document.add_paragraph(f'Сумма пополнений вклада  {self.amount_of_additions} ₽')
        document.add_heading('Результаты расчёта', level=1)
        document.add_paragraph(f'Сумма в конце срока  {round(compute_list[0][1])} ₽')
        document.add_paragraph(f'Доход  {round(compute_list[1][1])} ₽')
        document.add_heading('График начислений', level=1)

        table = document.add_table(rows=1, cols=5, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        column_widths = [Cm(1.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0)]
        for idx, width in enumerate(column_widths):
            table.columns[idx].width = width

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Дата начисления'
        hdr_cells[2].text = 'Начислено процентов'
        hdr_cells[3].text = 'Пополнение вклада'
        hdr_cells[4].text = 'Остаток вклада'

        for i, row in enumerate(data_lst):
            table.add_row()
            for j, val in enumerate(row):
                cell = table.cell(i+1, j)
                if j > 1:
                    cell.text = str(val[1]) + ' ₽'
                else:
                    cell.text = str(val[1])

        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        document.add_page_break()
        if copy_num == 0:
            document.save('deposit.docx')
        else:
            document.save(f'deposit{copy_num}.docx')

